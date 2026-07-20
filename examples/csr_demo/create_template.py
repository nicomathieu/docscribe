"""Generate a realistic (but entirely fictional) CSR template DOCX.

Run once to create examples/csr_demo/templates/csr_template.docx

    python examples/csr_demo/create_template.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def bold_run(para, text: str, size: int = 11) -> None:
    r = para.add_run(text)
    r.bold = True
    r.font.size = Pt(size)


def plain_run(para, text: str = "", size: int = 11) -> None:
    r = para.add_run(text)
    r.bold = False
    r.font.size = Pt(size)


def add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Add a 2-column key-value table with bold labels."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, placeholder) in enumerate(rows):
        label_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]

        label_para = label_cell.paragraphs[0]
        r = label_para.add_run(label)
        r.bold = True
        r.font.size = Pt(10)

        value_para = value_cell.paragraphs[0]
        vr = value_para.add_run(placeholder)
        vr.bold = False
        vr.font.size = Pt(10)


def add_multifield_table(doc: Document, title: str, fields: list[tuple[str, str]]) -> None:
    """Add a single-column table where each cell has a bold label + blank value."""
    doc.add_paragraph(title, style="Heading 2")
    table = doc.add_table(rows=len(fields), cols=1)
    table.style = "Table Grid"
    for i, (label, hint) in enumerate(fields):
        para = table.rows[i].cells[0].paragraphs[0]
        bold_run(para, f"{label}: ")
        plain_run(para, hint)


def main() -> None:
    out_dir = Path(__file__).parent / "templates"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / "csr_template.docx"

    doc = Document()

    # ── Title page ──────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("CLINICAL STUDY REPORT")
    r.bold = True
    r.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run("[FICTIONAL COMPOUND — FOR DEMONSTRATION ONLY]")
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # ── Section 1: Study Identification ─────────────────────────────────────
    add_section_heading(doc, "1. Study Identification", level=1)
    add_kv_table(doc, [
        ("Study Title:", ""),
        ("Protocol Number:", ""),
        ("EudraCT / NCT Number:", ""),
        ("Phase:", ""),
        ("Sponsor:", ""),
        ("Coordinating Investigator:", ""),
        ("Study Start Date:", ""),
        ("Study End Date:", ""),
        ("Report Date:", ""),
    ])

    doc.add_paragraph()

    # ── Section 2: Study Synopsis ────────────────────────────────────────────
    add_section_heading(doc, "2. Synopsis", level=1)
    add_multifield_table(doc, "", [
        ("Objectives", "[primary and secondary objectives]"),
        ("Design", "[study design summary]"),
        ("Population", "[inclusion/exclusion criteria summary]"),
        ("Treatment", "[IMP dosing and comparator]"),
        ("Primary Endpoint", "[primary efficacy endpoint]"),
        ("Key Secondary Endpoints", "[list of secondary endpoints]"),
        ("Statistical Methods", "[brief statistical analysis plan]"),
        ("Number of Subjects", "[planned and analysed]"),
    ])

    doc.add_paragraph()

    # ── Section 3: Investigational Product ───────────────────────────────────
    add_section_heading(doc, "3. Investigational Medicinal Product", level=1)
    add_kv_table(doc, [
        ("INN / Generic Name:", ""),
        ("Proprietary Name:", ""),
        ("ATC Code:", ""),
        ("Formulation:", ""),
        ("Route of Administration:", ""),
        ("Dose and Regimen:", ""),
        ("Manufacturer:", ""),
        ("Batch / Lot Number:", ""),
    ])

    doc.add_paragraph()

    # ── Section 4: Efficacy Results ───────────────────────────────────────────
    add_section_heading(doc, "4. Efficacy Results", level=1)

    add_section_heading(doc, "4.1 Primary Endpoint", level=2)
    table4_1 = doc.add_table(rows=4, cols=3)
    table4_1.style = "Table Grid"
    headers = ["Parameter", "Treatment Arm", "Control Arm"]
    for i, h in enumerate(headers):
        para = table4_1.rows[0].cells[i].paragraphs[0]
        r = para.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
    for row_label, row_idx in [("n (analysed)", 1), ("Mean ± SD", 2), ("p-value", 3)]:
        para = table4_1.rows[row_idx].cells[0].paragraphs[0]
        r = para.add_run(row_label)
        r.bold = True
        r.font.size = Pt(10)

    add_section_heading(doc, "4.2 Secondary Endpoints Summary", level=2)
    add_multifield_table(doc, "", [
        ("Secondary Endpoint 1", "[result]"),
        ("Secondary Endpoint 2", "[result]"),
        ("Secondary Endpoint 3", "[result]"),
    ])

    doc.add_paragraph()

    # ── Section 5: Safety Results ─────────────────────────────────────────────
    add_section_heading(doc, "5. Safety Results", level=1)
    add_kv_table(doc, [
        ("Total Subjects Exposed:", ""),
        ("Subjects with ≥1 AE:", ""),
        ("Subjects with ≥1 SAE:", ""),
        ("Discontinuations due to AE:", ""),
        ("Deaths:", ""),
    ])

    add_section_heading(doc, "5.1 Most Common Adverse Events (≥5%)", level=2)
    table5 = doc.add_table(rows=6, cols=4)
    table5.style = "Table Grid"
    headers5 = ["Adverse Event (MedDRA PT)", "Treatment Arm n (%)", "Control Arm n (%)", "All n (%)"]
    for i, h in enumerate(headers5):
        para = table5.rows[0].cells[i].paragraphs[0]
        r = para.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for row_idx in range(1, 6):
        para = table5.rows[row_idx].cells[0].paragraphs[0]
        para.add_run(f"AE {row_idx}").font.size = Pt(9)

    doc.add_paragraph()

    # ── Section 6: Conclusions ────────────────────────────────────────────────
    add_section_heading(doc, "6. Conclusions", level=1)
    add_multifield_table(doc, "", [
        ("Efficacy Conclusion", "[overall efficacy narrative]"),
        ("Safety Conclusion", "[overall safety narrative]"),
        ("Benefit-Risk Assessment", "[benefit-risk statement]"),
    ])

    doc.add_paragraph()

    # ── Section 7: Signatures ─────────────────────────────────────────────────
    add_section_heading(doc, "7. Signatures", level=1)
    add_kv_table(doc, [
        ("Medical Officer:", ""),
        ("Biostatistician:", ""),
        ("Clinical Pharmacologist:", ""),
        ("Date of Sign-off:", ""),
    ])

    doc.save(str(out_path))
    print(f"Template created: {out_path}")


if __name__ == "__main__":
    main()
