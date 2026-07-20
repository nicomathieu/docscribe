"""DOCX ↔ AST conversion using python-docx."""

from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from docfill.ast.models import (
    DocumentAST,
    HeadingElement,
    ParagraphElement,
    TableCellElement,
    TextRun,
    compute_hash,
)


def _para_text(para) -> str:
    return "".join(r.text for r in para.runs)


def _extract_runs(para) -> list[TextRun]:
    runs = []
    for r in para.runs:
        if not r.text:
            continue
        font_size = None
        if r.font.size:
            font_size = r.font.size.pt
        runs.append(
            TextRun(
                text=r.text,
                bold=bool(r.bold),
                italic=bool(r.italic),
                underline=r.underline if r.underline is not None else None,
                font_name=r.font.name or None,
                font_size_pt=font_size,
            )
        )
    return runs


def _cell_text(cell) -> str:
    return "\n".join(_para_text(p) for p in cell.paragraphs if _para_text(p))


def _cell_runs(cell) -> list[TextRun]:
    all_runs: list[TextRun] = []
    for i, para in enumerate(cell.paragraphs):
        if i > 0 and all_runs:
            all_runs.append(TextRun(text="\n", bold=False, italic=False))
        all_runs.extend(_extract_runs(para))
    return all_runs


def build_ast(file_path: str | Path) -> DocumentAST:
    """Parse a DOCX file and return its AST."""
    path = Path(file_path)
    doc = Document(str(path))

    with open(path, "rb") as f:
        file_hash = compute_hash(f.read())

    file_id = re.sub(r"[^a-z0-9-]", "", path.stem.lower().replace("_", "-").replace(" ", "-"))
    elements: list[Any] = []
    order = 0
    block_index = 0

    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "p":
            from docx.text.paragraph import Paragraph
            para = Paragraph(block, doc)
            text = _para_text(para)
            runs = _extract_runs(para)

            style_name = (para.style.name or "").lower() if para.style else ""
            if style_name.startswith("heading"):
                try:
                    level = int(style_name.split()[-1])
                except ValueError:
                    level = 1
                elements.append(
                    HeadingElement(
                        element_id=f"heading-{block_index}",
                        block_index=block_index,
                        level=level,
                        text=text,
                        order=order,
                        runs=runs,
                    )
                )
            else:
                elements.append(
                    ParagraphElement(
                        element_id=f"para-{block_index}",
                        block_index=block_index,
                        text=text,
                        order=order,
                        runs=runs,
                    )
                )
            block_index += 1
            order += 1

        elif tag == "tbl":
            from docx.table import Table
            table = Table(block, doc)
            table_idx = sum(1 for e in elements if hasattr(e, "table_index"))

            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    text = _cell_text(cell)
                    runs = _cell_runs(cell)
                    content_hash = compute_hash(text.encode())
                    elements.append(
                        TableCellElement(
                            element_id=f"cell-t{table_idx}-r{row_idx}-c{col_idx}",
                            table_index=table_idx,
                            row_index=row_idx,
                            col_index=col_idx,
                            text=text,
                            order=order,
                            runs=runs,
                            content_hash=content_hash,
                        )
                    )
                    order += 1
            block_index += 1

    return DocumentAST(
        file_id=file_id,
        filename=path.name,
        elements=elements,
        total_elements=len(elements),
    )


def _find_cell(doc: Document, table_index: int, row_index: int, col_index: int):
    tables = doc.tables
    if table_index >= len(tables):
        raise ValueError(f"Table {table_index} not found (doc has {len(tables)} tables)")
    table = tables[table_index]
    rows = table.rows
    if row_index >= len(rows):
        raise ValueError(f"Row {row_index} out of range")
    cells = rows[row_index].cells
    if col_index >= len(cells):
        raise ValueError(f"Col {col_index} out of range")
    return cells[col_index]


def apply_cell_edit(file_path: str | Path, cell_element: TableCellElement) -> str:
    """Write *cell_element.runs* (or .text) back into the DOCX at the matching cell."""
    path = Path(file_path)
    doc = Document(str(path))

    cell = _find_cell(doc, cell_element.table_index, cell_element.row_index, cell_element.col_index)

    # Clear existing paragraphs (keep at least one)
    for para in cell.paragraphs[1:]:
        p = para._element
        p.getparent().remove(p)

    first_para = cell.paragraphs[0]
    # Clear runs in first paragraph
    for child in list(first_para._element):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            first_para._element.remove(child)

    runs = cell_element.runs or []
    current_para = first_para

    for run in runs:
        if "\n" in run.text:
            parts = run.text.split("\n")
            for i, part in enumerate(parts):
                if i > 0:
                    # Add new paragraph
                    new_para = _add_paragraph_after(cell, current_para)
                    current_para = new_para
                if part:
                    _write_run(current_para, part, run)
        else:
            if run.text:
                _write_run(current_para, run.text, run)

    doc.save(str(path))
    return f"Applied edit to cell t{cell_element.table_index}-r{cell_element.row_index}-c{cell_element.col_index}"


def _add_paragraph_after(cell, current_para):
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    current_para._element.addnext(p)
    from docx.text.paragraph import Paragraph
    return Paragraph(p, cell._tc)


def _write_run(para, text: str, run: TextRun) -> None:
    r = para.add_run(text)
    r.bold = run.bold
    r.italic = run.italic
    if run.underline is not None:
        r.underline = run.underline
    if run.font_name:
        r.font.name = run.font_name
    if run.font_size_pt:
        r.font.size = Pt(run.font_size_pt)


def apply_heading_edit(file_path: str | Path, heading: HeadingElement) -> str:
    """Update a heading paragraph in the DOCX."""
    path = Path(file_path)
    doc = Document(str(path))

    block_index = heading.block_index
    blocks = [b for b in doc.element.body]
    para_blocks = [b for b in blocks if b.tag.split("}")[-1] == "p"]

    if block_index >= len(para_blocks):
        raise ValueError(f"Heading block_index {block_index} out of range")

    from docx.text.paragraph import Paragraph
    para = Paragraph(para_blocks[block_index], doc)
    for child in list(para._element):
        if child.tag.endswith("}r"):
            para._element.remove(child)

    r = para.add_run(heading.text)
    doc.save(str(path))
    return f"Applied heading edit to block {block_index}"
